import docx
import sys
import os
import re

files = [
    "examples/raquisitos fm.docx",
    "examples/REQUISITOS TÉCNICOS de cONDIÇÕES DE USO DE RADIOFREQUÊNCIAS PARA OS SERVIÇOS DE RADIODIFUSÃO DE SONS E IMAGENS E DE RETRANSMISSÃO DE TELEVISÃO.docx"
]

keywords = ["Classe", "ERP", "Altura", "Nível", "Contorno", "Proteção"]

for filename in files:
    if not os.path.exists(filename):
        continue
        
    print(f"\n\n========== ANALYZING: {os.path.basename(filename)} ==========\n")
    try:
        doc = docx.Document(filename)
        
        # Check paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if "FM" in text and any(k in text for k in keywords):
                print(f"[P{i}] {text}")
                
        # Check tables - this is where class definitions usually are
        print("\n--- RELEVANT TABLES ---\n")
        for t_idx, table in enumerate(doc.tables):
            # Check if table header or content looks like FM classes
            is_relevant = False
            rows_data = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_text)
                if any(x in " ".join(row_text) for x in ["Classe", "ERP", "Máxima", "Referência"]):
                    is_relevant = True
            
            if is_relevant:
                print(f"Table {t_idx}:")
                for row in rows_data:
                    print(" | ".join(row))
                print("-" * 40)

    except Exception as e:
        print(f"Error: {e}")
