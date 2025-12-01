import docx
import sys
import os

filename = "examples/REQUISITOS TÉCNICOS de cONDIÇÕES DE USO DE RADIOFREQUÊNCIAS PARA OS SERVIÇOS DE RADIODIFUSÃO DE SONS E IMAGENS E DE RETRANSMISSÃO DE TELEVISÃO.docx"

try:
    if not os.path.exists(filename):
        print(f"File not found: {filename}")
        # Try to find it with wildcard or similar if name is tricky
        for f in os.listdir("examples"):
            if "REQUISITOS" in f:
                filename = os.path.join("examples", f)
                print(f"Found file: {filename}")
                break

    doc = docx.Document(filename)
    print(f"Reading: {filename}")
    
    # Print headings and paragraphs that might contain "Classe" or "Nível" or "dBuV"
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (any(k in text for k in ["Classe", "Nível", "dB", "ERP", "Potência"]) or len(text) < 100):
            print(text)
            
    print("\n--- TABLES ---\n")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            # Filter for relevant tables
            if any("Classe" in c or "ERP" in c or "Nível" in c for c in row_text):
                 print(" | ".join(row_text))

except Exception as e:
    print(f"Error reading file: {e}")
