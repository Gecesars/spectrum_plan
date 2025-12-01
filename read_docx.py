import docx
import sys

try:
    doc = docx.Document("examples/raquisitos fm.docx")
    print("File found: raquisitos fm.docx")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    # Also check tables as requirements are often in tables
    print("\n--- TABLES ---\n")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_text))
            
except Exception as e:
    print(f"Error reading file: {e}")
