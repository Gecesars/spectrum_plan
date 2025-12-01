import docx
import sys
import os

filename = "examples/raquisitos fm.docx"

try:
    doc = docx.Document(filename)
    print(f"Reading: {filename}")
    
    # Search for specific FM class names
    fm_classes = ["A1", "A2", "A3", "A4", "B1", "B2", "E1", "E2", "E3"]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(c in text for c in fm_classes) and ("Classe" in text or "kW" in text):
            print(f"[P{i}] {text}")

    print("\n--- TABLES WITH FM CLASSES ---\n")
    for t_idx, table in enumerate(doc.tables):
        rows_data = []
        is_fm_table = False
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_text)
            row_str = " ".join(row_text)
            if any(c in row_str for c in fm_classes) and ("kW" in row_str or "W" in row_str):
                is_fm_table = True
        
        if is_fm_table:
            print(f"Table {t_idx}:")
            for row in rows_data:
                print(" | ".join(row))
            print("-" * 40)

except Exception as e:
    print(f"Error reading file: {e}")
